from __future__ import annotations

"""Starter utilities for rebuilding a section in a Word document.

Copy this module into a project workspace and adapt the content builder to the
target PDD template. It is intentionally generic rather than project-specific.
"""

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


def iter_block_items(parent: DocumentObject):
    body = parent._body._element
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def find_heading_paragraph(doc: Document, text: str) -> Paragraph:
    for item in iter_block_items(doc):
        if isinstance(item, Paragraph) and item.text.strip() == text:
            return item
    raise RuntimeError(f"Heading not found: {text}")


def remove_content_between(doc: Document, start_heading: str, end_heading: str) -> tuple[Paragraph, Paragraph]:
    start = find_heading_paragraph(doc, start_heading)
    end = find_heading_paragraph(doc, end_heading)

    body = doc._body._element
    found = False
    for child in list(body.iterchildren()):
        if child is start._p:
            found = True
            continue
        if child is end._p:
            break
        if found:
            body.remove(child)
    return start, end


def insert_paragraph_before(anchor: Paragraph, style: str | None = None) -> Paragraph:
    doc = anchor._parent
    paragraph = doc.add_paragraph(style=style)
    paragraph._p.getparent().remove(paragraph._p)
    anchor._p.addprevious(paragraph._p)
    return paragraph


def insert_table_before(anchor: Paragraph, rows: int, cols: int) -> Table:
    doc = anchor._parent
    table = doc.add_table(rows=rows, cols=cols)
    table._tbl.getparent().remove(table._tbl)
    anchor._p.addprevious(table._tbl)
    return table


def add_note_paragraph(anchor: Paragraph, text: str) -> Paragraph:
    p = insert_paragraph_before(anchor)
    p.add_run("Note: ").bold = True
    p.add_run(text).italic = True
    fmt = p.paragraph_format
    fmt.space_after = Pt(4)
    return p


def set_cell_width(cell: _Cell, inches: float) -> None:
    width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width)))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = width


def shade_cell(cell: _Cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_height(row, points: float) -> None:
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    row.height = Pt(points)


def format_table(table: Table, widths: list[float], *, body_font_size: float = 8.5) -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:
            repeat_table_header(row)
            set_row_height(row, 22)
        else:
            set_row_height(row, 18)

        for col_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_width(cell, widths[min(col_idx, len(widths) - 1)])
            if row_idx == 0:
                shade_cell(cell, "D9EAF7")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(1)
                for run in paragraph.runs:
                    run.font.size = Pt(9 if row_idx == 0 else body_font_size)
                    if row_idx == 0:
                        run.bold = True


def build_section(docx_path: Path, start_heading: str, end_heading: str) -> None:
    """Adapt this function per project."""
    doc = Document(str(docx_path))
    _, anchor = remove_content_between(doc, start_heading, end_heading)

    insert_paragraph_before(anchor, style="Heading2").add_run("Detailed TO BE Process Map")
    add_note_paragraph(anchor, "Replace this scaffold with component-specific content.")

    table = insert_table_before(anchor, 1, 3)
    table.cell(0, 0).text = "Column A"
    table.cell(0, 1).text = "Column B"
    table.cell(0, 2).text = "Column C"
    row = table.add_row()
    row.cells[0].text = "Example"
    row.cells[1].text = "Adapt this module for the project"
    row.cells[2].text = "Keep headings and numbering native to Word"
    format_table(table, [1.2, 2.4, 2.2])

    doc.save(str(docx_path))


if __name__ == "__main__":
    raise SystemExit("Import this module into a project-specific script and adapt build_section().")
